from __future__ import annotations

import subprocess
import tempfile
from collections.abc import Generator, Iterable, Iterator
from pathlib import Path
from typing import Self, final
from unittest.mock import patch

import pytest
from rich.console import Console

from quarry.api import (
    DeleteCollectionRequest,
    DeregisterAccepted,
    DeregisterRequest,
    RegisterRequest,
    RegistrationInfo,
    RegistrationList,
    RetainedCollection,
    TaskAccepted,
)
from quarry.client import QuarryClient
from quarry.config import Settings
from quarry.db import Database
from quarry.db.storage import get_db
from quarry.db_pointer import SELECTION
from quarry.ingestion.backends import clear_caches
from quarry.scratch_paths import ScratchGuard
from quarry.types import LanceDB
from tests.fakes import FakeEmbeddingBackend
from tests.hermetic_env import ENV, ProductionTreeGuard
from tests.inproc_daemon import InProcessDaemon

FIXTURES_DIR = Path(__file__).parent / "fixtures"

_REPO_ROOT = Path(__file__).resolve().parent.parent


def _pytest_tmp_base() -> Path:
    """Return a temp base for ``tmp_path`` that the watcher/indexer never refuses.

    Tests build simulated project roots via ``tmp_path``.  The workspace exports
    ``TMPDIR=$PWD/.tmp``, so pytest would place those roots under the repo's own
    ``.tmp`` scratch — which the guard now refuses (DES-045).  ``<repo>/.pytest-
    work`` sidesteps that in CI.  A dev worktree can itself live under a parent
    repo's ``.tmp`` (this one is ``<quarry>/.tmp/watch-wt``); the guard then also
    refuses ``<repo>/.pytest-work``, so fall back to a home-cache base that no
    repo's scratch encloses.  Either way the base is a directory the guard
    permits, so a test's project root stays indexable.
    """
    candidate = _REPO_ROOT / ".pytest-work"
    if ScratchGuard().refuses_root(candidate):
        return Path.home() / ".cache" / "quarry-pytest"
    return candidate


def pytest_configure(config: pytest.Config) -> None:
    """Point pytest's temp base at a guard-permitted directory (DES-045).

    Overriding ``tempfile.tempdir`` (not ``--basetemp``) preserves pytest's
    ``pytest-of-<user>/pytest-<n>`` rotation and its concurrency safety.
    """
    del config
    base = _pytest_tmp_base()
    base.mkdir(parents=True, exist_ok=True)
    tempfile.tempdir = str(base)


@pytest.fixture(scope="session")
def asgi_daemon(
    tmp_path_factory: pytest.TempPathFactory,
) -> Iterator[QuarryClient]:
    """A hermetic in-process daemon: real handlers over ASGI, no socket, no ONNX.

    Session-scoped so the app and its (empty) tmp database build once. Yields a
    :class:`QuarryClient` wired to the real daemon app; a CLI/MCP test patches
    ``TargetResolver.connect`` to return this client and exercises the real
    ``CLI → client → daemon → LanceDB`` request path with NO live ``quarryd``
    running — the guarantee the "daemon-mandatory tests must be hermetic" rule
    demands (verifiable with the daemon STOPPED).
    """
    data_dir = tmp_path_factory.mktemp("asgi-daemon")
    daemon = InProcessDaemon(data_dir)
    with daemon.client() as client:
        yield client


class GitSandbox:
    """A throwaway git repository the ratchet tests drive end-to-end.

    The merge-base ratchets (``tools/oo_ratchet``, ``tools/coupling``,
    ``tools/suppression``) read the *base-commit* baseline blob via
    ``git show <base>:<file>``, so they can only be exercised faithfully against
    a real commit graph — a mocked ``GitRepo`` drifts from the code it stands in
    for (the recurring remote/local-divergence bug class). This sandbox commits
    real trees so the tests score against genuine base blobs.
    """

    _root: Path

    def __new__(cls, root: Path) -> Self:
        self = super().__new__(cls)
        self._root = root
        self._git("init", "-q")
        self._git("config", "user.email", "ratchet@test.local")
        self._git("config", "user.name", "Ratchet Test")
        self._git("config", "commit.gpgsign", "false")
        return self

    @property
    def root(self) -> Path:
        """Return the repository root directory."""
        return self._root

    def write(self, relpath: str, text: str) -> Path:
        """Write ``text`` to ``relpath`` under the repo, creating parents."""
        path = self._root / relpath
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(text)
        return path

    def commit(self, message: str) -> str:
        """Stage everything and commit, returning the full commit hash."""
        self._git("add", "-A")
        self._git("commit", "-q", "-m", message)
        return self._git("rev-parse", "HEAD").strip()

    def run_git(self, *args: str) -> str:
        """Run an arbitrary git command in the repo and return its stdout.

        Lets a test seed refs the ratchets read — e.g. an
        ``refs/remotes/origin/main`` pointer for the coupling adoption-tripwire
        path — without a real remote.
        """
        return self._git(*args)

    def _git(self, *args: str) -> str:
        result = subprocess.run(
            ["git", *args],
            cwd=self._root,
            capture_output=True,
            text=True,
            check=True,
        )
        return result.stdout


@pytest.fixture()
def git_sandbox(tmp_path: Path) -> GitSandbox:
    """Return a fresh :class:`GitSandbox` rooted at ``tmp_path``."""
    return GitSandbox(tmp_path)


@final
class FakeRegistryClient:
    """In-memory stand-in for the daemon registry surface enable/disable use.

    Structurally satisfies ``quarry.enable.RegistryClient`` and records
    register/deregister/delete_collection calls so a test can assert the client
    dispatched exactly the daemon operations the design requires — never touching
    a local ``SyncRegistry``.  Construct with ``(collection, directory)`` pairs to
    seed a covering ``RegistrationList``.
    """

    __slots__ = (
        "_chunk_collections",
        "_delete_error",
        "_deleted",
        "_deregistered",
        "_registered",
        "_regs",
        "_retained",
    )

    _regs: list[RegistrationInfo]
    _registered: list[RegisterRequest]
    _deregistered: list[DeregisterRequest]
    _deleted: list[str]
    # Collections archived by a keep-data deregister: {collection: original_dir},
    # keyed by collection name with its origin directory as the value — mirroring
    # the real retained_collections (PK ``collection``, ``original_directory``
    # column) so the client sees archived names AND their origin (for re-adopt).
    _retained: dict[str, str]
    # Chunk-bearing collection names the daemon reports on the wire, so the
    # client-side picker avoids a name that already holds another project's chunks.
    _chunk_collections: list[str]
    # When set, delete_collection raises it — models a rejected captures purge.
    _delete_error: Exception | None

    def __new__(
        cls,
        registrations: Iterable[tuple[str, Path]] = (),
        *,
        delete_error: Exception | None = None,
        chunk_collections: Iterable[str] = (),
    ) -> Self:
        self = super().__new__(cls)
        self._regs = [
            RegistrationInfo(
                collection=col,
                directory=str(Path(directory).resolve()),
                registered_at="2026-01-01",
            )
            for col, directory in registrations
        ]
        self._registered = []
        self._deregistered = []
        self._deleted = []
        self._retained = {}
        self._chunk_collections = list(chunk_collections)
        self._delete_error = delete_error
        return self

    def list_registrations(self) -> RegistrationList:
        retained = [
            RetainedCollection(collection=col, original_directory=directory)
            for col, directory in sorted(self._retained.items())
        ]
        return RegistrationList(
            total_registrations=len(self._regs),
            registrations=list(self._regs),
            retained=retained,
            chunk_collections=list(self._chunk_collections),
        )

    def register(self, req: RegisterRequest) -> TaskAccepted:
        self._registered.append(req)
        self._retained.pop(req.collection, None)  # re-adopting clears the archive
        self._regs.append(
            RegistrationInfo(
                collection=req.collection,
                directory=req.directory,
                registered_at="2026-01-02",
            )
        )
        return TaskAccepted(task_id="t")

    def deregister(self, req: DeregisterRequest) -> DeregisterAccepted:
        self._deregistered.append(req)
        origin = next(
            (r.directory for r in self._regs if r.collection == req.collection), ""
        )
        before = len(self._regs)
        self._regs = [r for r in self._regs if r.collection != req.collection]
        if req.keep_data and before > len(self._regs):
            self._retained[req.collection] = origin  # kept → archived, with origin
        return DeregisterAccepted(task_id="t", removed=before - len(self._regs))

    def delete_collection(self, req: DeleteCollectionRequest) -> TaskAccepted:
        if self._delete_error is not None:
            raise self._delete_error
        self._deleted.append(req.name)
        return TaskAccepted(task_id="t")

    @property
    def registered(self) -> list[RegisterRequest]:
        return self._registered

    @property
    def deregistered(self) -> list[DeregisterRequest]:
        return self._deregistered

    @property
    def deleted(self) -> list[str]:
        return self._deleted

    @property
    def collections(self) -> list[str]:
        return [r.collection for r in self._regs]


# Environment variables that .envrc / shell may set and that pydantic-settings
# would otherwise silently inject into Settings instances created in tests.
_QUARRY_ENV_VARS = (
    "EMBEDDING_DIMENSION",
    "EMBEDDING_MODEL",
    "LANCEDB_PATH",
    "LOG_PATH",
    "QUARRY_API_KEY",
    "QUARRY_PROVIDER",
    "REGISTRY_PATH",
)
# QUARRY_ROOT is deliberately absent: ``tests.hermetic_env`` points it at the
# session home, and stripping it would send every Settings() back to the real
# ``~/.punt-labs/quarry/data`` -- away from isolation, not toward it.


@pytest.fixture(autouse=True)
def _no_remote_config() -> Generator[None]:
    """Prevent tests from reading the real mcp-proxy config on disk.

    ``TargetResolver.resolve`` reads ``read_proxy_config`` for the stored-remote
    tier; returning ``{}`` keeps resolution on the loopback default unless a test
    patches the resolver itself.
    """
    with patch("quarry.client.resolver.read_proxy_config", return_value={}):
        yield


@pytest.fixture(autouse=True)
def _isolate_from_env(monkeypatch: pytest.MonkeyPatch) -> None:
    """Strip quarry-related env vars so tests get deterministic Settings defaults.

    Without this, .envrc exports leak into every Settings() call, causing
    spurious failures when the test assumes the code-level default.
    """
    for var in _QUARRY_ENV_VARS:
        monkeypatch.delenv(var, raising=False)


@pytest.fixture()
def real_embedding_factory() -> None:
    """Request this to keep the real embedding factory in place.

    For the factory's own unit tests, which assert on what
    ``get_embedding_backend`` builds and supply their own ONNX-session mocks.
    Distinct from ``@pytest.mark.embedding``: this opts out of the fake but
    stays in the default run and still may not load a real model.
    """
    return


@pytest.fixture(autouse=True)
def _fake_embedding_backend(request: pytest.FixtureRequest) -> Generator[None]:
    """Install the shared fake at the embedding factory, unless opted out.

    Three symbols are replaced, not two, and the third is the one that does the
    work.  Both ``quarry.ingestion.streaming`` and ``quarry.http_resources``
    bind their factory with ``from quarry.ingestion.backends import ...`` at
    module scope, so patching the factory module alone leaves those bound names
    on the real implementation.  ``OnnxEmbeddingBackend`` is imported inside
    ``new_embedding_backend`` and therefore resolves per call, which makes it
    the one point every route -- cached, fresh, current, or added tomorrow --
    passes through.

    A test carrying ``@pytest.mark.embedding`` gets the real factory. The cache
    is cleared on both edges so a real backend built by such a test cannot be
    handed to the next one.
    """
    keeps_real_factory = (
        request.node.get_closest_marker("embedding") is not None
        or "real_embedding_factory" in request.fixturenames
    )
    if keeps_real_factory:
        clear_caches()
        yield
        clear_caches()
        return

    clear_caches()
    with (
        patch(
            "quarry.ingestion.backends.get_embedding_backend",
            FakeEmbeddingBackend.for_settings,
        ),
        patch("quarry.ingestion.backends.new_embedding_backend", FakeEmbeddingBackend),
        patch("quarry.embeddings.OnnxEmbeddingBackend", FakeEmbeddingBackend),
    ):
        yield
    clear_caches()


@pytest.fixture(autouse=True)
def _forbid_real_model_load(request: pytest.FixtureRequest) -> Generator[None]:
    """Fail loudly on an unmarked test that reaches a real ONNX session.

    The fake prevents the load; this proves the prevention holds. Without it a
    refactor that moves the factory turns "410 MB was loaded" back into a silent
    cost that leaves the test passing.
    """
    if request.node.get_closest_marker("embedding") is not None:
        yield
        return

    def _refuse(*args: object, **kwargs: object) -> object:
        del args, kwargs
        msg = (
            f"{request.node.nodeid} loaded a real ONNX model. Mark it with "
            "@pytest.mark.embedding if that is intended (it costs 410 MB and "
            "is deselected by default), or route it through the fake backend."
        )
        raise AssertionError(msg)

    with patch("onnxruntime.InferenceSession", _refuse):
        yield


@pytest.fixture(autouse=True)
def _reset_database_selection() -> Generator[None]:
    """Leave no database selection behind for the next test.

    ``SELECTION`` stands for a process-global choice, and its persisted half is
    a real file under the session home. Without this, a test that selects a
    database silently changes what its successors resolve to.

    The session pointer is resolved at *setup*, before the test can relocate the
    root. A path resolved only at teardown depends on whether ``monkeypatch``
    has already restored ``QUARRY_ROOT``, which pytest does not order against
    this fixture — so it could look under the wrong root and leave the real
    pointer in place. Both paths are removed, since a test may have written
    under either.
    """
    session_pointer = SELECTION.path
    yield
    SELECTION.override("")
    for pointer in {session_pointer, SELECTION.path}:
        pointer.unlink(missing_ok=True)


@pytest.fixture(autouse=True)
def _guard_production_tree() -> Generator[None]:
    """Fail the test that writes to the operator's real quarry tree.

    Runs unconditionally, including in CI where ``HOME`` differs but the same
    three paths must stay untouched.  Three stats per test measured at 25.9 ms
    across the whole suite.
    """
    guard = ProductionTreeGuard(ENV.real_tree)
    yield
    if breaches := guard.changed():
        pytest.fail("wrote to the production quarry tree:\n" + "\n".join(breaches))


@pytest.fixture(autouse=True)
def _force_no_color(monkeypatch: pytest.MonkeyPatch) -> None:
    """Force quarry's rich consoles to emit plain, ANSI-free CLI output.

    ``CliRunner`` captures output through a pipe, so rich would normally
    disable color on its own. But each ``Console(stderr=True)`` in
    ``src/quarry`` is built with ``force_terminal=None`` and holds a live
    ``os.environ`` reference: rich re-reads ``FORCE_COLOR`` at *render* time,
    not at construction. So a color-forcing ambient env makes these consoles
    emit ANSI escapes into the captured output at print time, breaking the
    plain substring assertions. Deleting the env var at runtime does not help
    — rich re-reads it on the next render. Replacing each console with one
    pinned to ``force_terminal=False, no_color=True`` defeats that render-time
    re-read, making the gate color-deterministic regardless of the shell.

    Every ``rich.console.Console`` holds a live reference to ``os.environ`` and
    reads ``NO_COLOR`` at *render* time, so setting it here makes every console
    (including the one ``CliPlumbing`` injected into each command group) emit
    plain output — no per-object patching needed.  Typer is not patched: every
    ``typer.Typer`` app sets ``rich_markup_mode=None``, so it renders through
    click's plain formatter, never ``rich_utils``.
    """
    monkeypatch.setenv("NO_COLOR", "1")
    monkeypatch.delenv("FORCE_COLOR", raising=False)
    monkeypatch.setattr(
        "quarry.__main__.err_console",
        Console(stderr=True, no_color=True, force_terminal=False),
    )


@pytest.fixture(scope="session")
def embedding_model_name() -> str:
    return "Snowflake/snowflake-arctic-embed-m-v1.5"


@pytest.fixture(scope="session")
def _warm_embedding_model(embedding_model_name: str) -> None:
    """Load ONNX embedding model once per session."""
    from quarry.ingestion.backends import get_embedding_backend

    settings = Settings(embedding_model=embedding_model_name)
    get_embedding_backend(settings).embed_texts(["warm up"])


@pytest.fixture()
def integration_settings(
    embedding_model_name: str, _warm_embedding_model: None
) -> Settings:
    """Settings with real embedding model."""
    return Settings(embedding_model=embedding_model_name)


@pytest.fixture()
def lance_db(tmp_path: Path) -> LanceDB:
    return get_db(tmp_path / "db")


@pytest.fixture()
def database(lance_db: LanceDB) -> Database:
    """Database facade wrapping the test ``lance_db`` connection."""
    return Database(lance_db)


@pytest.fixture(scope="session")
def pdf_fixture(tmp_path_factory: pytest.TempPathFactory) -> Path:
    """Generate 2-page PDF with embedded text via PyMuPDF."""
    import fitz

    tmp = tmp_path_factory.mktemp("fixtures")
    path = tmp / "test-document.pdf"

    doc = fitz.open()

    page1 = doc.new_page(width=612, height=792)
    page1.insert_text(
        (72, 72),
        (
            "Software Engineering Principles\n\n"
            "Software engineering applies systematic, disciplined approaches\n"
            "to the development, operation, and maintenance of software.\n"
            "Key practices include version control, automated testing,\n"
            "continuous integration, and code review."
        ),
        fontsize=12,
    )

    page2 = doc.new_page(width=612, height=792)
    page2.insert_text(
        (72, 72),
        (
            "Marine Biology Overview\n\n"
            "Marine biology studies organisms in the ocean and other\n"
            "saltwater environments. Coral reefs support approximately\n"
            "25 percent of all marine species despite covering less than\n"
            "one percent of the ocean floor."
        ),
        fontsize=12,
    )

    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture(scope="session")
def docx_fixture(tmp_path_factory: pytest.TempPathFactory) -> Path:
    """Generate DOCX with heading styles via python-docx."""
    from docx import Document

    tmp = tmp_path_factory.mktemp("fixtures")
    path = tmp / "test-document.docx"

    doc = Document()
    doc.add_heading("Distributed Systems", level=1)
    doc.add_paragraph(
        "Distributed systems coordinate multiple computers to achieve"
        " a common goal. Key challenges include consensus, fault tolerance,"
        " and network partitioning. The CAP theorem states that a distributed"
        " system cannot simultaneously guarantee consistency, availability,"
        " and partition tolerance."
    )

    doc.add_heading("Operating Systems", level=1)
    doc.add_paragraph(
        "Operating systems manage hardware resources and provide services"
        " for application software. Process scheduling, memory management,"
        " and file systems are core responsibilities. Modern operating"
        " systems use virtual memory to give each process an isolated"
        " address space."
    )

    doc.save(str(path))
    return path


@pytest.fixture(scope="session")
def png_fixture(tmp_path_factory: pytest.TempPathFactory) -> Path:
    """Generate PNG with text via PIL."""
    from PIL import Image, ImageDraw

    tmp = tmp_path_factory.mktemp("fixtures")
    path = tmp / "test-image.png"

    img = Image.new("RGB", (400, 100), color="white")
    draw = ImageDraw.Draw(img)
    draw.text((10, 40), "Hello OCR", fill="black")
    img.save(str(path))
    return path
